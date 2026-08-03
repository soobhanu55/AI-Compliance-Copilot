"""Drafts structured compliance documentation (e.g. AI Act Annex IV skeleton).

build_annex_iv_sections() is the single source of truth for section content — both the .docx
export and the JSON endpoint used by the "Drafted Docs" frontend page render from it, so they
never drift out of sync with each other.

Section text is template-generated from the company profile + gap-report findings, not an LLM
call — no LLM API key is configured in this environment. It's real, deterministic content
(not placeholder lorem ipsum), just not natural-language drafting.
"""

from pathlib import Path
from typing import TypedDict

from docx import Document

from app.models.schemas import ClauseAssessment, CompanyProfile, GapReport


class AnnexIVSection(TypedDict):
    number: int
    title: str
    citation: str
    body: str


SECTION_TITLES = [
    "General description of the system and intended purpose",
    "Detailed description of elements and development process",
    "Monitoring, functioning, and control information",
    "Risk management system description",
    "Changes made through the system lifecycle",
    "Standards and specifications applied",
    "EU declaration of conformity",
    "Post-market monitoring plan",
]

# Keywords used to pick the most relevant assessment to cite for each section — a light
# heuristic, not a claim of legal precision. Falls back to a generic Annex IV chip.
SECTION_KEYWORDS = [
    ["scope", "definitions", "classification"],
    ["technical documentation", "data and data governance", "record-keeping"],
    ["human oversight", "transparency", "accuracy"],
    ["risk management"],
    ["corrective actions", "post-market"],
    ["quality management", "standards"],
    ["conformity", "declaration"],
    ["post-market monitoring", "reporting of serious incidents"],
]


def _find_citation(assessments: list[ClauseAssessment], keywords: list[str]) -> str:
    for assessment in assessments:
        article_lower = assessment.article.lower()
        if any(kw in article_lower for kw in keywords):
            return f"{assessment.article}"
    return "Annex IV"


def _ai_system_summary(profile: CompanyProfile) -> str:
    return profile.ai_system_descriptions[0] if profile.ai_system_descriptions else "the AI system in question"


def build_annex_iv_sections(profile: CompanyProfile, report: GapReport) -> list[AnnexIVSection]:
    system = _ai_system_summary(profile)
    applicable = [a for a in report.assessments if a.verdict == "applicable"]
    needs_review = [a for a in report.assessments if a.verdict == "needs_human_review"]

    bodies = [
        f"The system is {system}, deployed within {profile.name}'s operations "
        f"({profile.sector}, {profile.employee_count} employees). Its intended purpose is "
        f"described in the company profile" + (f": {profile.notes}" if profile.notes else "."),
        f"Development process details are not yet documented. {len(applicable)} clause(s) were "
        f"classified as applicable by the proof-of-concept compliance classifier and require "
        f"a documented technical description here before this section can be considered complete.",
        "Monitoring, functioning, and control information — including human-oversight measures "
        "and accuracy/robustness testing results — is not yet documented for this system.",
        f"Risk management documentation is pending. {len(needs_review)} clause(s) across the "
        f"retrieved regulation set were flagged as needing human review and may bear on the "
        f"scope of the risk management system required here.",
        "No lifecycle changes have been logged yet for this system.",
        "Applicable harmonised standards and common specifications have not yet been identified "
        "for this system.",
        "An EU declaration of conformity has not yet been drafted or issued for this system.",
        "A post-market monitoring plan has not yet been established for this system.",
    ]

    sections: list[AnnexIVSection] = []
    for i, (title, keywords, body) in enumerate(zip(SECTION_TITLES, SECTION_KEYWORDS, bodies), start=1):
        sections.append(
            AnnexIVSection(
                number=i,
                title=title,
                citation=_find_citation(report.assessments, keywords),
                body=body,
            )
        )
    return sections


def draft_annex_iv(profile: CompanyProfile, report: GapReport, output_path: Path) -> Path:
    sections = build_annex_iv_sections(profile, report)

    doc = Document()
    doc.add_heading(f"AI Act Annex IV — Technical Documentation ({profile.name})", level=0)
    doc.add_paragraph(
        "DRAFT — auto-generated skeleton. Review with qualified counsel before submission."
    )

    for section in sections:
        doc.add_heading(f"{section['number']}. {section['title']} (§ {section['citation']})", level=1)
        doc.add_paragraph(section["body"])

    doc.add_heading("Relevant gap-report findings", level=1)
    for assessment in report.assessments:
        doc.add_paragraph(
            f"{assessment.regulation} {assessment.article}: {assessment.verdict} "
            f"({assessment.rationale})"
        )

    doc.save(output_path)
    return output_path
